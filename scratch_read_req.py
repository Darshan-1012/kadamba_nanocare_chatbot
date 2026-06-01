from docx import Document
doc = Document(r'app\summary_requirement.docx')
print('=== FULL PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] {p.text[:300]}')
print('\n=== ALL TABLES ===')
for i, table in enumerate(doc.tables):
    print(f'\n--- Table {i+1} ({len(table.rows)}r x {len(table.columns)}c) ---')
    for r, row in enumerate(table.rows):
        cells = [c.text.strip()[:100] for c in row.cells]
        print(f'  R{r}: ' + ' | '.join(cells))
